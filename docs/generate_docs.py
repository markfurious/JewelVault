"""
Generate comprehensive Word documentation for the Inventory/Jewelry ERP System.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def create_document():
    """Create the main document with all sections."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Inventory / Jewelry ERP System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete System Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph()
    doc.add_paragraph(f'Generated: March 31, 2026')
    doc.add_paragraph('Version: 1.0.0')

    # Table of Contents placeholder
    doc.add_page_break()

    # ========== SECTION 1: EXECUTIVE SUMMARY ==========
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('1.1 System Overview', level=2)
    doc.add_paragraph(
        'This document provides comprehensive documentation for the Inventory/Jewelry ERP System, '
        'a production-ready enterprise resource planning solution designed specifically for diamond '
        'and jewelry businesses. The system features an item-based inventory tracking model where '
        'each SKU represents exactly ONE physical item, tracked by status rather than quantity.'
    )

    doc.add_heading('1.2 Key Features', level=2)
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('Item-based inventory tracking (AVAILABLE/SOLD/RESERVED status)')
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('Product management with jewelry-specific attributes')
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('Sales processing with automatic inventory status updates')
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('Smart reorder suggestions based on sales velocity')
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('Bulk product upload via Excel')
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('User authentication and role-based access control')
    features = doc.add_paragraph(style='List Bullet')
    features.add_run('Comprehensive audit trail for all inventory transactions')

    doc.add_heading('1.3 Technology Stack', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Layer'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'

    tech_data = [
        ('Frontend', 'React 18 + Vite', 'UI rendering and state management'),
        ('Grid Component', 'AG Grid v35', 'Tabular data display'),
        ('Backend Framework', 'FastAPI', 'RESTful API server'),
        ('Validation', 'Pydantic', 'Request/response schema validation'),
        ('ORM', 'SQLAlchemy 2.0', 'Database object mapping'),
        ('Database', 'PostgreSQL', 'Persistent data storage'),
        ('Migrations', 'Alembic', 'Database schema versioning'),
        ('Authentication', 'JWT + bcrypt', 'Secure user authentication'),
    ]

    for layer, tech, purpose in tech_data:
        row = table.add_row()
        row.cells[0].text = layer
        row.cells[1].text = tech
        row.cells[2].text = purpose

    # ========== SECTION 2: SYSTEM ARCHITECTURE ==========
    doc.add_page_break()
    doc.add_heading('2. System Architecture', level=1)

    doc.add_heading('2.1 High-Level Architecture', level=2)
    doc.add_paragraph(
        'The system follows a layered architecture with clear separation of concerns:'
    )

    doc.add_paragraph('┌─────────────────────────────────────────────────────────┐\n'
                      '│                    USER BROWSER                         │\n'
                      '│               http://localhost:3002                     │\n'
                      '└────────────────────┬────────────────────────────────────┘\n'
                      '                     │ HTTP/HTTPS\n'
                      '                     ▼\n'
                      '┌─────────────────────────────────────────────────────────┐\n'
                      '│              REACT FRONTEND (Vite)                      │\n'
                      '│  Pages | Components | Services | Context                │\n'
                      '└────────────────────┬────────────────────────────────────┘\n'
                      '                     │ Proxy: /api/* → localhost:8000\n'
                      '                     ▼\n'
                      '┌─────────────────────────────────────────────────────────┐\n'
                      '│             FASTAPI BACKEND (uvicorn)                   │\n'
                      '│  API Routers | Services | Models | Schemas              │\n'
                      '└────────────────────┬────────────────────────────────────┘\n'
                      '                     │ SQLAlchemy\n'
                      '                     ▼\n'
                      '┌─────────────────────────────────────────────────────────┐\n'
                      '│              POSTGRESQL DATABASE                        │\n'
                      '│  products | inventory | sales | users | transactions    │\n'
                      '└─────────────────────────────────────────────────────────┘',
                      style='No Spacing')

    doc.add_heading('2.2 Component Responsibilities', level=2)

    doc.add_heading('Frontend Layer', level=3)
    doc.add_paragraph(
        'The React frontend is responsible for:', style='List Bullet')
    doc.add_paragraph('Rendering user interfaces for all business operations', style='List Bullet')
    doc.add_paragraph('Managing local state and form data', style='List Bullet')
    doc.add_paragraph('Displaying tabular data using AG Grid', style='List Bullet')
    doc.add_paragraph('Handling user authentication and session management', style='List Bullet')
    doc.add_paragraph('Making HTTP requests to the backend API', style='List Bullet')

    doc.add_heading('Backend API Layer', level=3)
    doc.add_paragraph(
        'The FastAPI backend handles:', style='List Bullet')
    doc.add_paragraph('RESTful endpoint routing', style='List Bullet')
    doc.add_paragraph('Request validation using Pydantic schemas', style='List Bullet')
    doc.add_paragraph('Business logic execution via service layer', style='List Bullet')
    doc.add_paragraph('Database operations through SQLAlchemy ORM', style='List Bullet')
    doc.add_paragraph('Exception handling and error responses', style='List Bullet')

    doc.add_heading('Service Layer', level=3)
    doc.add_paragraph(
        'Services encapsulate business logic:', style='List Bullet')
    doc.add_paragraph('ProductService: Product CRUD, SKU generation', style='List Bullet')
    doc.add_paragraph('InventoryService: Status changes, transaction logging', style='List Bullet')
    doc.add_paragraph('SaleService: Sale creation, cancellation, validation', style='List Bullet')
    doc.add_paragraph('AnalyticsService: Reorder suggestions, sales velocity', style='List Bullet')
    doc.add_paragraph('AuthService: JWT authentication, user management', style='List Bullet')

    # ========== SECTION 3: BACKEND STRUCTURE ==========
    doc.add_page_break()
    doc.add_heading('3. Backend Structure', level=1)

    doc.add_heading('3.1 Directory Structure', level=2)
    doc.add_paragraph(
        'backend/\n'
        '├── app/\n'
        '│   ├── __init__.py\n'
        '│   ├── main.py                    # FastAPI app factory\n'
        '│   │\n'
        '│   ├── api/                       # HTTP Interface Layer\n'
        '│   │   ├── __init__.py\n'
        '│   │   ├── dependencies.py        # get_db(), auth dependencies\n'
        '│   │   └── v1/\n'
        '│   │       ├── products.py        # CRUD + bulk upload\n'
        '│   │       ├── inventory.py       # Status adjustments\n'
        '│   │       ├── sales.py           # Sale transactions\n'
        '│   │       ├── stock.py           # Combined grid endpoint\n'
        '│   │       ├── analytics.py       # Reorder suggestions\n'
        '│   │       └── auth.py            # User management\n'
        '│   │\n'
        '│   ├── models/                    # SQLAlchemy ORM Models\n'
        '│   │   ├── product.py             # Product table\n'
        '│   │   ├── inventory.py           # Inventory + Transactions\n'
        '│   │   ├── sale.py                # Sales + SaleItems\n'
        '│   │   ├── user.py                # User authentication\n'
        '│   │   └── reorder.py             # Reorder rules\n'
        '│   │\n'
        '│   ├── schemas/                   # Pydantic Validation\n'
        '│   │   ├── product.py             # ProductCreate, ProductUpdate\n'
        '│   │   ├── inventory.py           # InventoryResponse\n'
        '│   │   ├── sale.py                # SaleCreate, SaleItemCreate\n'
        '│   │   └── auth.py                # Login, User schemas\n'
        '│   │\n'
        '│   ├── services/                  # Business Logic\n'
        '│   │   ├── product_service.py     # Product operations\n'
        '│   │   ├── inventory_service.py   # Status management\n'
        '│   │   ├── sale_service.py        # Sale processing\n'
        '│   │   ├── analytics_service.py   # Business intelligence\n'
        '│   │   └── auth_service.py        # Authentication\n'
        '│   │\n'
        '│   ├── db/                        # Database Configuration\n'
        '│   │   └── base.py                # Engine, session, Base\n'
        '│   │\n'
        '│   ├── core/                      # Core Configuration\n'
        '│   │   └── config.py              # Settings, environment\n'
        '│   │\n'
        '│   └── utils/                     # Utilities\n'
        '│       ├── exceptions.py          # Custom exceptions\n'
        '│       └── sku_generator.py       # SKU generation\n'
        '│\n'
        '├── alembic/                       # Database Migrations\n'
        '│   └── versions/\n'
        '│       ├── 001_initial_migration.py\n'
        '│       ├── 002_add_reorder_rules.py\n'
        '│       ├── 003_add_sales.py\n'
        '│       └── 004_convert_to_item_based.py\n'
        '│\n'
        '└── requirements.txt',
        style='No Spacing')

    doc.add_heading('3.2 Service Layer Details', level=2)

    service_table = doc.add_table(rows=1, cols=3)
    service_table.style = 'Table Grid'
    headers = service_table.rows[0].cells
    headers[0].text = 'Service'
    headers[1].text = 'Key Methods'
    headers[2].text = 'Responsibility'

    services = [
        ('ProductService',
         'create(), update(), delete(), list_all(), get_by_id()',
         'Product CRUD operations, SKU generation, validation'),
        ('ProductBulkService',
         'process_upload()',
         'Excel file parsing, bulk product creation'),
        ('InventoryService',
         'mark_sold(), mark_available(), reserve(), release_reservation(), adjust_status()',
         'Item status management, transaction logging'),
        ('SaleService',
         'create(), cancel(), list_all(), get_sales_by_product()',
         'Sale creation with validation, cancellation'),
        ('AnalyticsService',
         'get_reorder_suggestions(), get_sales_velocity()',
         'Business intelligence, predictive analytics'),
    ]

    for name, methods, responsibility in services:
        row = service_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = methods
        row.cells[2].text = responsibility

    # ========== SECTION 4: DATABASE SCHEMA ==========
    doc.add_page_break()
    doc.add_heading('4. Database Schema', level=1)

    doc.add_heading('4.1 Entity Relationship Diagram', level=2)
    doc.add_paragraph(
        'Product (1) ────── (1) Inventory\n'
        '   │                    │\n'
        '   │                    │ (1:N)\n'
        '   │                    ▼\n'
        '   │              InventoryTransaction\n'
        '   │\n'
        '   │ (1:N)\n'
        '   ▼\n'
        'SaleItem (N) ────── (1) Sale',
        style='No Spacing')

    doc.add_heading('4.2 Table Definitions', level=2)

    # Products table
    doc.add_heading('products', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Column'
    table.rows[0].cells[1].text = 'Type'
    table.rows[0].cells[2].text = 'Nullable'
    table.rows[0].cells[3].text = 'Description'

    products_cols = [
        ('id', 'UUID', 'No', 'Primary key'),
        ('sku', 'VARCHAR(20)', 'No', 'Stock keeping unit (unique)'),
        ('name', 'VARCHAR(255)', 'No', 'Product name'),
        ('description', 'TEXT', 'Yes', 'Product description'),
        ('category', 'VARCHAR(100)', 'Yes', 'Product category'),
        ('sub_category', 'VARCHAR(100)', 'Yes', 'Product sub-category'),
        ('style_number', 'VARCHAR(50)', 'Yes', 'Jewelry style number'),
        ('st_carat', 'NUMERIC(10,4)', 'Yes', 'Stone carat weight'),
        ('product_weight', 'NUMERIC(10,4)', 'Yes', 'Product weight in grams'),
        ('gold_purity', 'VARCHAR(20)', 'Yes', 'Gold purity (14K, 18K, etc.)'),
        ('certified', 'BOOLEAN', 'Yes', 'Certification flag'),
        ('cost_price', 'NUMERIC(12,2)', 'Yes', 'Cost price'),
        ('retail_price', 'NUMERIC(12,2)', 'Yes', 'Retail price'),
        ('wholesale_price', 'NUMERIC(12,2)', 'Yes', 'Wholesale price'),
        ('online_price', 'NUMERIC(12,2)', 'Yes', 'Online price'),
        ('is_active', 'BOOLEAN', 'No', 'Active status'),
        ('attributes', 'JSONB', 'Yes', 'Extensible attributes'),
        ('default_reorder_threshold', 'NUMERIC(10,2)', 'No', 'Reorder threshold'),
        ('created_at', 'TIMESTAMP', 'No', 'Creation timestamp'),
        ('updated_at', 'TIMESTAMP', 'No', 'Last update timestamp'),
    ]

    for col in products_cols:
        row = table.add_row()
        for i, val in enumerate(col):
            row.cells[i].text = val

    # Inventory table
    doc.add_heading('inventory', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Column'
    table.rows[0].cells[1].text = 'Type'
    table.rows[0].cells[2].text = 'Nullable'
    table.rows[0].cells[3].text = 'Description'

    inventory_cols = [
        ('id', 'UUID', 'No', 'Primary key'),
        ('product_id', 'UUID', 'No', 'Foreign key to products (unique)'),
        ('status', 'VARCHAR(20)', 'No', 'Item status: AVAILABLE/SOLD/RESERVED'),
        ('location', 'VARCHAR(255)', 'Yes', 'Storage location'),
        ('warehouse_code', 'VARCHAR(50)', 'Yes', 'Warehouse identifier'),
        ('created_at', 'TIMESTAMP', 'No', 'Creation timestamp'),
        ('updated_at', 'TIMESTAMP', 'No', 'Last update timestamp'),
    ]

    for col in inventory_cols:
        row = table.add_row()
        for i, val in enumerate(col):
            row.cells[i].text = val

    # Sales table
    doc.add_heading('sales', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Column'
    table.rows[0].cells[1].text = 'Type'
    table.rows[0].cells[2].text = 'Nullable'
    table.rows[0].cells[3].text = 'Description'

    sales_cols = [
        ('id', 'UUID', 'No', 'Primary key'),
        ('sale_number', 'VARCHAR(50)', 'No', 'Unique sale identifier'),
        ('customer_name', 'VARCHAR(255)', 'Yes', 'Customer name'),
        ('customer_email', 'VARCHAR(255)', 'Yes', 'Customer email'),
        ('customer_phone', 'VARCHAR(50)', 'Yes', 'Customer phone'),
        ('sale_type', 'VARCHAR(20)', 'No', 'RETAIL/WHOLESALE/ONLINE'),
        ('subtotal', 'NUMERIC(12,2)', 'No', 'Subtotal amount'),
        ('tax_amount', 'NUMERIC(12,2)', 'No', 'Tax amount'),
        ('discount_amount', 'NUMERIC(12,2)', 'No', 'Discount amount'),
        ('total_amount', 'NUMERIC(12,2)', 'No', 'Total amount'),
        ('payment_method', 'VARCHAR(50)', 'Yes', 'CASH/CARD/TRANSFER'),
        ('payment_status', 'VARCHAR(20)', 'Yes', 'Payment status'),
        ('status', 'VARCHAR(20)', 'No', 'COMPLETED/CANCELLED/REFUNDED'),
        ('notes', 'TEXT', 'Yes', 'Sale notes'),
        ('sale_date', 'TIMESTAMP', 'No', 'Sale timestamp'),
        ('created_at', 'TIMESTAMP', 'No', 'Creation timestamp'),
        ('updated_at', 'TIMESTAMP', 'No', 'Last update timestamp'),
    ]

    for col in sales_cols:
        row = table.add_row()
        for i, val in enumerate(col):
            row.cells[i].text = val

    # Inventory Transactions table
    doc.add_heading('inventory_transactions', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Column'
    table.rows[0].cells[1].text = 'Type'
    table.rows[0].cells[2].text = 'Nullable'
    table.rows[0].cells[3].text = 'Description'

    trans_cols = [
        ('id', 'UUID', 'No', 'Primary key'),
        ('inventory_id', 'UUID', 'No', 'Foreign key to inventory'),
        ('transaction_type', 'VARCHAR(50)', 'No', 'SALE/RESTOCK/ADJUSTMENT/etc.'),
        ('status_before', 'VARCHAR(20)', 'No', 'Status before transaction'),
        ('status_after', 'VARCHAR(20)', 'No', 'Status after transaction'),
        ('reference_id', 'UUID', 'Yes', 'Related entity ID (e.g., sale_id)'),
        ('reference_type', 'VARCHAR(50)', 'Yes', 'Related entity type'),
        ('notes', 'VARCHAR(500)', 'Yes', 'Transaction notes'),
        ('performed_by', 'VARCHAR(100)', 'Yes', 'User who performed action'),
        ('created_at', 'TIMESTAMP', 'No', 'Transaction timestamp'),
    ]

    for col in trans_cols:
        row = table.add_row()
        for i, val in enumerate(col):
            row.cells[i].text = val

    # ========== SECTION 5: API ENDPOINTS ==========
    doc.add_page_break()
    doc.add_heading('5. API Endpoints', level=1)

    doc.add_heading('5.1 Products API', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Endpoint'
    table.rows[0].cells[1].text = 'Method'
    table.rows[0].cells[2].text = 'Description'

    endpoints = [
        ('/api/v1/products', 'POST', 'Create new product'),
        ('/api/v1/products', 'GET', 'List products (paginated)'),
        ('/api/v1/products/{id}', 'GET', 'Get product by ID'),
        ('/api/v1/products/sku/{sku}', 'GET', 'Get product by SKU'),
        ('/api/v1/products/{id}', 'PUT', 'Update product'),
        ('/api/v1/products/{id}', 'DELETE', 'Soft delete product'),
        ('/api/v1/products/bulk-upload', 'POST', 'Bulk upload from Excel'),
    ]

    for endpoint, method, desc in endpoints:
        row = table.add_row()
        row.cells[0].text = endpoint
        row.cells[1].text = method
        row.cells[2].text = desc

    doc.add_heading('5.2 Inventory API', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Endpoint'
    table.rows[0].cells[1].text = 'Method'
    table.rows[0].cells[2].text = 'Description'

    endpoints = [
        ('/api/v1/inventory', 'GET', 'List inventory (paginated)'),
        ('/api/v1/inventory/{product_id}', 'GET', 'Get inventory for product'),
        ('/api/v1/inventory/{product_id}', 'PUT', 'Update inventory'),
        ('/api/v1/inventory/{product_id}/adjust', 'POST', 'Adjust status'),
        ('/api/v1/inventory/{product_id}/transactions', 'GET', 'Get transaction history'),
    ]

    for i, (endpoint, method, desc) in enumerate(endpoints):
        row = table.add_row()
        row.cells[0].text = endpoint
        row.cells[1].text = method
        row.cells[2].text = desc

    doc.add_heading('5.3 Sales API', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Endpoint'
    table.rows[0].cells[1].text = 'Method'
    table.rows[0].cells[2].text = 'Description'

    endpoints = [
        ('/api/v1/sales', 'GET', 'List sales (paginated)'),
        ('/api/v1/sales', 'POST', 'Create new sale'),
        ('/api/v1/sales/{id}', 'GET', 'Get sale by ID'),
        ('/api/v1/sales/{id}/cancel', 'POST', 'Cancel sale'),
    ]

    for endpoint, method, desc in endpoints:
        row = table.add_row()
        row.cells[0].text = endpoint
        row.cells[1].text = method
        row.cells[2].text = desc

    doc.add_heading('5.4 Stock API (Combined)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Endpoint'
    table.rows[0].cells[1].text = 'Method'
    table.rows[0].cells[2].text = 'Description'

    endpoints = [
        ('/api/v1/stock', 'GET', 'Get stock grid data (products + inventory)'),
        ('/api/v1/stock/summary', 'GET', 'Get stock summary statistics'),
        ('/api/v1/stock/{product_id}', 'GET', 'Get detailed stock info'),
    ]

    for endpoint, method, desc in endpoints:
        row = table.add_row()
        row.cells[0].text = endpoint
        row.cells[1].text = method
        row.cells[2].text = desc

    doc.add_heading('5.5 Analytics API', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Endpoint'
    table.rows[0].cells[1].text = 'Method'
    table.rows[0].cells[2].text = 'Description'

    endpoints = [
        ('/api/v1/analytics/reorder-suggestions', 'GET', 'Get reorder suggestions'),
        ('/api/v1/analytics/sales-velocity', 'GET', 'Get sales velocity data'),
        ('/api/v1/analytics/inventory/summary', 'GET', 'Get inventory summary'),
        ('/api/v1/analytics/top-products', 'GET', 'Get top selling products'),
        ('/api/v1/analytics/query', 'POST', 'Execute custom analytics query'),
    ]

    for endpoint, method, desc in endpoints:
        row = table.add_row()
        row.cells[0].text = endpoint
        row.cells[1].text = method
        row.cells[2].text = desc

    # ========== SECTION 6: DATA FLOWS ==========
    doc.add_page_break()
    doc.add_heading('6. Data Flow Diagrams', level=1)

    doc.add_heading('6.1 Creating a Product', level=2)
    doc.add_paragraph(
        '1. User fills product form in ProductsPage\n'
        '2. Frontend calls productsApi.create(payload)\n'
        '3. Backend validates ProductCreate schema\n'
        '4. ProductService.create() executes:\n'
        '   - Generate/validate SKU (format: SI00001)\n'
        '   - Create Product record\n'
        '   - Create Inventory record (status=AVAILABLE)\n'
        '   - Create ReorderRule with default threshold\n'
        '5. Database commits all records\n'
        '6. Response returned to frontend\n'
        '7. AG Grid refreshes with new data',
        style='No Spacing')

    doc.add_heading('6.2 Selling an Item (Item-Based Model)', level=2)
    doc.add_paragraph(
        '1. User navigates to CreateSalePage\n'
        '2. Frontend loads AVAILABLE items only:\n'
        '   GET /api/v1/stock?available_only=true\n'
        '3. User selects products, enters quantities\n'
        '4. User submits sale with customer info\n'
        '5. Backend validates each item:\n'
        '   - Product must exist\n'
        '   - Inventory must exist\n'
        '   - Status MUST be AVAILABLE (validation!)\n'
        '6. SaleService.create() executes:\n'
        '   - Create Sale record\n'
        '   - Create SaleItem records\n'
        '   - Call InventoryService.mark_sold() for each item\n'
        '7. InventoryService creates transaction:\n'
        '   - status_before: AVAILABLE\n'
        '   - status_after: SOLD\n'
        '   - transaction_type: SALE\n'
        '8. Response returned, frontend navigates to sales list',
        style='No Spacing')

    doc.add_heading('6.3 Status Adjustment Workflow', level=2)
    doc.add_paragraph(
        '1. User double-clicks row in InventoryPage\n'
        '2. Adjust form opens with current status\n'
        '3. User selects new status (AVAILABLE/SOLD/RESERVED)\n'
        '4. User enters notes (reason for change)\n'
        '5. Frontend calls inventoryApi.adjust():\n'
        '   POST /api/v1/inventory/{product_id}/adjust\n'
        '   {status, transaction_type, notes}\n'
        '6. InventoryService.adjust_status():\n'
        '   - Validates new status\n'
        '   - Creates InventoryTransaction record\n'
        '   - Updates inventory.status\n'
        '7. Grid refreshes showing new status',
        style='No Spacing')

    # ========== SECTION 7: FRONTEND STRUCTURE ==========
    doc.add_page_break()
    doc.add_heading('7. Frontend Structure', level=1)

    doc.add_heading('7.1 Directory Structure', level=2)
    doc.add_paragraph(
        'frontend/\n'
        '├── src/\n'
        '│   ├── main.jsx                   # React entry point\n'
        '│   ├── App.jsx                    # Main app with routing\n'
        '│   │\n'
        '│   ├── pages/                     # Page Components\n'
        '│   │   ├── DashboardPage.jsx      # Dashboard overview\n'
        '│   │   ├── StockPage.jsx          # Stock grid view\n'
        '│   │   ├── ProductsPage.jsx       # Product management\n'
        '│   │   ├── InventoryPage.jsx      # Inventory adjustments\n'
        '│   │   ├── SalesPage.jsx          # Sales history\n'
        '│   │   ├── CreateSalePage.jsx     # New sale form\n'
        '│   │   ├── ReorderSuggestionsPage.jsx  # Smart suggestions\n'
        '│   │   ├── LoginPage.jsx          # Login form\n'
        '│   │   └── AdminUsersPage.jsx     # User management\n'
        '│   │\n'
        '│   ├── components/                # Reusable Components\n'
        '│   │   ├── Layout.jsx             # Main layout with sidebar\n'
        '│   │   └── BulkUploadModal.jsx    # Excel upload modal\n'
        '│   │\n'
        '│   ├── services/                  # API Services\n'
        '│   │   └── api.js                 # All API calls\n'
        '│   │\n'
        '│   └── context/                   # React Context\n'
        '│       └── AuthContext.jsx        # Authentication state\n'
        '│\n'
        '├── public/\n'
        '├── vite.config.js                 # Vite configuration\n'
        '└── package.json',
        style='No Spacing')

    doc.add_heading('7.2 Page Components', level=2)

    page_table = doc.add_table(rows=1, cols=3)
    page_table.style = 'Table Grid'
    page_table.rows[0].cells[0].text = 'Page'
    page_table.rows[0].cells[1].text = 'API Calls'
    page_table.rows[0].cells[2].text = 'Purpose'

    pages = [
        ('DashboardPage', 'stockApi, inventoryApi, salesApi',
         'Overview statistics and quick metrics'),
        ('StockPage', 'stockApi.list(), productsApi.bulkUpload()',
         'Main stock grid with AG Grid display'),
        ('ProductsPage', 'stockApi.list(), productsApi.create()',
         'Product creation and management'),
        ('InventoryPage', 'inventoryApi.list(), inventoryApi.adjust()',
         'Individual item status adjustments'),
        ('SalesPage', 'salesApi.list()',
         'Sales history with filtering'),
        ('CreateSalePage', 'stockApi.list(available_only=true), salesApi.create()',
         'New sale creation with product selection'),
        ('ReorderSuggestionsPage', 'analyticsApi.getReorderSuggestions()',
         'AI-ready reorder recommendations'),
        ('LoginPage', 'authApi.login()',
         'User authentication'),
        ('AdminUsersPage', 'authApi.getUsers(), authApi.toggleUserActive()',
         'User management and permissions'),
    ]

    for name, apis, purpose in pages:
        row = page_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = apis
        row.cells[2].text = purpose

    doc.add_heading('7.3 API Service Layer', level=2)
    doc.add_paragraph(
        'The api.js service provides a centralized interface for all HTTP requests:'
    )

    doc.add_paragraph(
        'Key features:\n'
        '• Centralized error handling\n'
        '• Automatic JSON parsing\n'
        '• HTTP status code extraction\n'
        '• Token-based authentication support\n'
        '• Query string parameter handling',
        style='List Bullet')

    doc.add_paragraph(
        'API modules exported:\n'
        '• productsApi - Product CRUD operations\n'
        '• inventoryApi - Inventory status management\n'
        '• salesApi - Sale transactions\n'
        '• analyticsApi - Business intelligence\n'
        '• stockApi - Combined product/inventory data\n'
        '• authApi - Authentication and user management',
        style='List Bullet')

    # ========== SECTION 8: ITEM-BASED MODEL ==========
    doc.add_page_break()
    doc.add_heading('8. Item-Based Inventory Model', level=1)

    doc.add_heading('8.1 Concept Overview', level=2)
    doc.add_paragraph(
        'The system uses an ITEM-BASED inventory model, which is fundamentally different from '
        'traditional quantity-based systems. This model is specifically designed for jewelry and '
        'high-value item businesses where each item is unique and needs individual tracking.'
    )

    doc.add_heading('Key Principles', level=3)
    doc.add_paragraph('Each SKU represents exactly ONE physical item', style='List Bullet')
    doc.add_paragraph('Items are tracked by STATUS, not quantity', style='List Bullet')
    doc.add_paragraph('Status values: AVAILABLE, SOLD, RESERVED', style='List Bullet')
    doc.add_paragraph('Cannot sell an item unless status is AVAILABLE', style='List Bullet')
    doc.add_paragraph('Every status change is logged in inventory_transactions', style='List Bullet')

    doc.add_heading('8.2 Comparison: Quantity-Based vs Item-Based', level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Aspect'
    table.rows[0].cells[1].text = 'Quantity-Based (Traditional)'
    table.rows[0].cells[2].text = 'Item-Based (Current System)'

    comparisons = [
        ('Tracking Unit', 'Quantity count (e.g., 10 units)', 'Individual item status'),
        ('Sale Effect', 'Reduces quantity by N', 'Changes status to SOLD'),
        ('Validation', 'Check if quantity > 0', 'Check if status == AVAILABLE'),
        ('Best For', 'Commodity products, bulk items', 'Unique items, jewelry, art'),
        ('Inventory Record', 'One record per product', 'One record per physical item'),
        ('Restocking', 'Add to quantity', 'Change status to AVAILABLE'),
    ]

    for aspect, qty_based, item_based in comparisons:
        row = table.add_row()
        row.cells[0].text = aspect
        row.cells[1].text = qty_based
        row.cells[2].text = item_based

    doc.add_heading('8.3 Status State Machine', level=2)
    doc.add_paragraph(
        '┌─────────────┐\n'
        '│  AVAILABLE  │ ← Initial state for all new items\n'
        '└──────┬──────┘\n'
        '       │\n'
        '       ├─────────────┬─────────────┐\n'
        '       │             │             │\n'
        '       ▼             ▼             │\n'
        '┌─────────────┐ ┌─────────────┐   │\n'
        '│    SOLD     │ │  RESERVED   │   │\n'
        '└──────┬──────┘ └──────┬──────┘   │\n'
        '       │               │          │\n'
        '       │               │          │\n'
        '       │               ▼          │\n'
        '       │         ┌─────────────┐  │\n'
        '       │         │  RELEASED   │  │\n'
        '       │         └──────┬──────┘  │\n'
        '       │                │         │\n'
        '       └────────────────┴─────────┘\n'
        '                (back to AVAILABLE)',
        style='No Spacing')

    doc.add_heading('Valid Transitions', level=3)
    doc.add_paragraph('AVAILABLE → SOLD (via sale)', style='List Bullet')
    doc.add_paragraph('AVAILABLE → RESERVED (via reservation)', style='List Bullet')
    doc.add_paragraph('RESERVED → AVAILABLE (release reservation)', style='List Bullet')
    doc.add_paragraph('RESERVED → SOLD (sell reserved item)', style='List Bullet')
    doc.add_paragraph('SOLD → AVAILABLE (cancel sale/restock)', style='List Bullet')
    doc.add_paragraph('Any → Any (manual adjustment with notes)', style='List Bullet')

    # ========== SECTION 9: AUTHENTICATION ==========
    doc.add_page_break()
    doc.add_heading('9. Authentication System', level=1)

    doc.add_heading('9.1 User Model', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Column'
    table.rows[0].cells[1].text = 'Type'
    table.rows[0].cells[2].text = 'Nullable'
    table.rows[0].cells[3].text = 'Description'

    user_cols = [
        ('id', 'UUID', 'No', 'Primary key'),
        ('username', 'VARCHAR(100)', 'No', 'Unique username'),
        ('email', 'VARCHAR(255)', 'No', 'User email'),
        ('hashed_password', 'VARCHAR(255)', 'No', 'bcrypt hashed password'),
        ('role', 'VARCHAR(50)', 'No', 'admin/staff/manager'),
        ('is_active', 'BOOLEAN', 'No', 'Active account flag'),
        ('is_locked', 'BOOLEAN', 'No', 'Locked account flag'),
        ('created_at', 'TIMESTAMP', 'No', 'Creation timestamp'),
        ('updated_at', 'TIMESTAMP', 'No', 'Last update timestamp'),
    ]

    for col in user_cols:
        row = table.add_row()
        for i, val in enumerate(col):
            row.cells[i].text = val

    doc.add_heading('9.2 Authentication Flow', level=2)
    doc.add_paragraph(
        '1. User enters credentials on LoginPage\n'
        '2. Frontend calls authApi.login(username, password)\n'
        '3. Backend validates credentials via AuthService\n'
        '4. If valid, JWT tokens generated:\n'
        '   - access_token (short-lived)\n'
        '   - refresh_token (long-lived)\n'
        '5. Tokens stored in localStorage\n'
        '6. Subsequent requests include Authorization header:\n'
        '   Bearer <access_token>\n'
        '7. Token refreshed automatically when expired',
        style='No Spacing')

    doc.add_heading('9.3 User Roles', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Role'
    table.rows[0].cells[1].text = 'Permissions'

    roles = [
        ('admin', 'Full system access, user management, all CRUD operations'),
        ('manager', 'Product/inventory/sales management, no user admin'),
        ('staff', 'View-only access, can create sales'),
    ]

    for role, perms in roles:
        row = table.add_row()
        row.cells[0].text = role
        row.cells[1].text = perms

    # ========== SECTION 10: BULK UPLOAD ==========
    doc.add_page_break()
    doc.add_heading('10. Bulk Upload Feature', level=1)

    doc.add_heading('10.1 Excel Template Format', level=2)
    doc.add_paragraph(
        'Required columns in the Excel file:'
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Column Name'
    table.rows[0].cells[1].text = 'Required'
    table.rows[0].cells[2].text = 'Format/Notes'

    columns = [
        ('SKU', 'Yes', 'Format: SI00001 (SI + 5 digits)'),
        ('Category', 'No', 'e.g., Rings, Necklaces, Bracelets'),
        ('Sub Category', 'No', 'Optional sub-category'),
        ('Style Number', 'No', 'Product style identifier'),
        ('ST Carat', 'No', 'Stone carat weight (decimal)'),
        ('Product wt', 'No', 'Product weight in grams'),
        ('Gold Purity', 'No', '10K, 14K, 18K, 22K, 24K'),
        ('Certified', 'No', 'TRUE/FALSE'),
        ('Wholesale Price', 'No', 'Decimal number'),
        ('Retail Price', 'Yes', 'Decimal number'),
        ('Online Price', 'No', 'Decimal number'),
    ]

    for col, req, fmt in columns:
        row = table.add_row()
        row.cells[0].text = col
        row.cells[1].text = req
        row.cells[2].text = fmt

    doc.add_heading('10.2 Upload Process', level=2)
    doc.add_paragraph(
        '1. User clicks "Bulk Upload" button\n'
        '2. Modal opens with file picker\n'
        '3. User selects Excel file (.xlsx only)\n'
        '4. User specifies initial quantity and location\n'
        '5. Frontend uploads file via FormData\n'
        '6. Backend validates ALL rows before insertion\n'
        '7. If ANY row fails, entire file is rejected\n'
        '8. On success, all products created with AVAILABLE status\n'
        '9. Grid refreshes showing new items',
        style='No Spacing')

    doc.add_heading('10.3 Validation Rules', level=2)
    doc.add_paragraph('SKU must match pattern ^SI\\d{5}$', style='List Bullet')
    doc.add_paragraph('SKU must be unique (not already in database)', style='List Bullet')
    doc.add_paragraph('Retail Price is required', style='List Bullet')
    doc.add_paragraph('Numeric fields must be valid numbers', style='List Bullet')
    doc.add_paragraph('Gold Purity must be valid (10K-24K)', style='List Bullet')

    # ========== SECTION 11: ANALYTICS ==========
    doc.add_page_break()
    doc.add_heading('11. Analytics & Reporting', level=1)

    doc.add_heading('11.1 Reorder Suggestions', level=2)
    doc.add_paragraph(
        'The system provides AI-ready reorder suggestions based on:'
    )
    doc.add_paragraph('Sales velocity (units sold per day)', style='List Bullet')
    doc.add_paragraph('Current stock status (AVAILABLE count)', style='List Bullet')
    doc.add_paragraph('Minimum threshold settings', style='List Bullet')
    doc.add_paragraph('Estimated days until stockout', style='List Bullet')
    doc.add_paragraph('Recommended reorder quantity', style='List Bullet')

    doc.add_heading('11.2 Sales Velocity Calculation', level=2)
    doc.add_paragraph(
        'Sales velocity is calculated as:\n\n'
        'velocity = total_units_sold / days_in_period\n\n'
        'The system analyzes completed sales over a configurable period '
        '(default: 30 days) to determine how fast each product sells.'
    )

    doc.add_heading('11.3 Inventory Summary', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Description'

    metrics = [
        ('Total Products', 'Count of all products in system'),
        ('Products with Inventory', 'Products that have inventory records'),
        ('Available Count', 'Items with status = AVAILABLE'),
        ('Sold Count', 'Items with status = SOLD'),
        ('Reserved Count', 'Items with status = RESERVED'),
        ('Total Inventory Value', 'Sum of cost_price for AVAILABLE items'),
        ('Category Breakdown', 'Product count by category'),
        ('Gold Purity Breakdown', 'Product count by purity'),
    ]

    for metric, desc in metrics:
        row = table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = desc

    # ========== SECTION 12: MIGRATIONS ==========
    doc.add_page_break()
    doc.add_heading('12. Database Migrations', level=1)

    doc.add_heading('12.1 Migration History', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Version'
    table.rows[0].cells[1].text = 'Description'
    table.rows[0].cells[2].text = 'Key Changes'

    migrations = [
        ('001', 'Initial Migration',
         'Created products, inventory, users tables'),
        ('002', 'Add Reorder Rules',
         'Added reorder_rules table for AI suggestions'),
        ('003', 'Add Sales Tables',
         'Added sales, sale_items, inventory_transactions'),
        ('004', 'Convert to Item-Based',
         'Removed quantity fields, added status tracking'),
    ]

    for ver, desc, changes in migrations:
        row = table.add_row()
        row.cells[0].text = ver
        row.cells[1].text = desc
        row.cells[2].text = changes

    doc.add_heading('12.2 Running Migrations', level=2)
    doc.add_paragraph(
        '# Check current migration version\n'
        'alembic current\n\n'
        '# Upgrade to latest\n'
        'alembic upgrade head\n\n'
        '# Downgrade one version\n'
        'alembic downgrade -1\n\n'
        '# Create new migration\n'
        'alembic revision -m "description"',
        style='No Spacing')

    # ========== SECTION 13: RISKS & RECOMMENDATIONS ==========
    doc.add_page_break()
    doc.add_heading('13. Risks and Recommendations', level=1)

    doc.add_heading('13.1 Current Risks', level=2)

    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = 'Table Grid'
    risk_table.rows[0].cells[0].text = 'Risk'
    risk_table.rows[0].cells[1].text = 'Impact'
    risk_table.rows[0].cells[2].text = 'Mitigation'

    risks = [
        ('No database failover', 'High - System outage if DB fails',
         'Implement read replicas, connection retry logic'),
        ('No caching layer', 'Medium - Slow performance at scale',
         'Add Redis for frequently accessed data'),
        ('Synchronous bulk upload', 'Medium - Timeout for large files',
         'Implement Celery for async processing'),
        ('No rate limiting', 'Medium - API abuse possible',
         'Add slowapi middleware'),
        ('Hardcoded tax rate', 'Low - Requires deployment to change',
         'Move to database configuration'),
        ('Limited frontend error handling', 'Medium - Poor UX on failures',
         'Add retry logic, error boundaries'),
    ]

    for risk, impact, mitigation in risks:
        row = risk_table.add_row()
        row.cells[0].text = risk
        row.cells[1].text = impact
        row.cells[2].text = mitigation

    doc.add_heading('13.2 Recommended Next Steps', level=2)
    doc.add_paragraph('Add comprehensive test suite (pytest, React Testing Library)', style='List Bullet')
    doc.add_paragraph('Implement Redis caching for grid endpoints', style='List Bullet')
    doc.add_paragraph('Add Celery for background job processing', style='List Bullet')
    doc.add_paragraph('Implement rate limiting middleware', style='List Bullet')
    doc.add_paragraph('Add audit logging for all user actions', style='List Bullet')
    doc.add_paragraph('Move configuration (tax, thresholds) to database', style='List Bullet')
    doc.add_paragraph('Add TypeScript for type safety', style='List Bullet')

    # ========== SECTION 14: QUICK REFERENCE ==========
    doc.add_page_break()
    doc.add_heading('14. Quick Reference', level=1)

    doc.add_heading('14.1 Common Operations', level=2)

    doc.add_heading('Start Backend Server', level=3)
    doc.add_paragraph(
        'cd backend\n'
        'source venv/bin/activate\n'
        'uvicorn app.main:app --reload --host 0.0.0.0 --port 8000',
        style='No Spacing')

    doc.add_heading('Start Frontend Dev Server', level=3)
    doc.add_paragraph(
        'cd frontend\n'
        'npm run dev',
        style='No Spacing')

    doc.add_heading('Run Database Migrations', level=3)
    doc.add_paragraph(
        'cd backend\n'
        'alembic upgrade head',
        style='No Spacing')

    doc.add_heading('14.2 Default Credentials', level=2)
    doc.add_paragraph(
        'Note: Change these immediately in production!',
        style='List Bullet')
    doc.add_paragraph(
        'Username: admin\n'
        'Password: admin123',
        style='No Spacing')

    doc.add_heading('14.3 API Documentation URLs', level=2)
    doc.add_paragraph('Swagger UI: http://localhost:8000/docs', style='List Bullet')
    doc.add_paragraph('ReDoc: http://localhost:8000/redoc', style='List Bullet')
    doc.add_paragraph('Health Check: http://localhost:8000/health', style='List Bullet')

    # Add footer
    doc.add_page_break()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph(
        'Inventory / Jewelry ERP System - Complete Documentation\n'
        'Version 1.0.0 | Generated March 31, 2026',
        style='No Spacing')

    return doc


if __name__ == '__main__':
    doc = create_document()
    doc.save('/Users/markfurious/claude-ws/inventory-system/docs/Inventory_ERP_Documentation.docx')
    print('Documentation generated successfully!')
